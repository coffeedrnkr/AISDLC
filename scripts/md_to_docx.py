import os
import sys
import re
import argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: 'python-docx' library is not installed.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

def parse_markdown_line(text):
    """Simple parser for bold and italic within text."""
    parts = []
    segments = text.split('**')
    for i, seg in enumerate(segments):
        is_bold = (i % 2 == 1)
        # Handle code snippets inside text too `code`
        # For professional look, maybe we skip code style in simple text for now
        # or implement it. Let's stick to Bold support.
        parts.append({'text': seg, 'bold': is_bold})
    return parts

def process_table(doc, table_lines):
    """Parses markdown table lines and creates a styled docx table."""
    if len(table_lines) < 2:
        # Not enough lines for a table, dump as text
        for l in table_lines:
            doc.add_paragraph(l)
        return

    # 1. Parse Header
    header_line = table_lines[0].strip()
    headers = [h.strip() for h in header_line.strip('|').split('|')]
    
    # 2. Check/Skip Separator
    # MD separator often looks like |---|---|
    start_row_idx = 1
    if len(table_lines) > 1 and set(table_lines[1].strip()) <= set("|-: "):
        start_row_idx = 2

    # 3. Parse Data
    data_rows = []
    for line in table_lines[start_row_idx:]:
        if not line.strip(): continue
        row_cells = [c.strip() for c in line.strip('|').split('|')]
        # Handle mismatch cols
        if len(row_cells) != len(headers):
             # Pad or truncate?
             # Simple pad
             while len(row_cells) < len(headers): row_cells.append("")
             row_cells = row_cells[:len(headers)]
        data_rows.append(row_cells)

    # 4. Create Table
    rows = len(data_rows) + 1
    cols = len(headers)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 5. Format Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        # Clear default param
        tc = hdr_cells[i]
        p = tc.paragraphs[0]
        p.clear()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        
        # Add Shading (Light Gray)
        shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
        tc._tc.get_or_add_tcPr().append(shading_elm)

    # 6. Fill Rows
    for r_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[r_idx+1].cells
        for c_idx, content in enumerate(row_data):
            # Parse bold in cells?
            p = row_cells[c_idx].paragraphs[0]
            p.clear()
            # Simple bold parsing for cells too
            parts = parse_markdown_line(content)
            for part in parts:
                r = p.add_run(part['text'])
                if part['bold']: r.bold = True
                r.font.size = Pt(10)

def markdown_to_docx(input_file, output_file):
    doc = Document()
    
    # Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    table_buffer = []

    for line in lines:
        stripped = line.strip()

        # 1. Code Blocks
        if stripped.startswith("```"):
            # Flush table
            if table_buffer:
                process_table(doc, table_buffer)
                table_buffer = []
            
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(line.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
            
        # 2. Table Handling
        # Standard GFM table check
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                process_table(doc, table_buffer)
                table_buffer = []

        # 3. Headers
        if stripped.startswith("#"):
            level = len(stripped.split()[0])
            content = stripped.lstrip("#").strip()
            if level <= 9:
                doc.add_heading(content, level=level)
            else:
                doc.add_paragraph(content, style='Body Text')
            continue

        # 4. Lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_markdown_line(content)
            p.clear() 
            for part in parts:
                run = p.add_run(part['text'])
                run.bold = part['bold']
            continue
            
        # 5. Normal Text
        if not stripped:
            continue
            
        p = doc.add_paragraph()
        parts = parse_markdown_line(stripped)
        for part in parts:
            run = p.add_run(part['text'])
            run.bold = part['bold']

    # Final flush
    if table_buffer:
        process_table(doc, table_buffer)

    doc.save(output_file)
    print(f"✅ Converted '{input_file}' to '{output_file}' with Professional Tables")

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX")
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", help="Output DOCX file", nargs='?')
    
    args = parser.parse_args()
    
    if not args.output:
        args.output = os.path.splitext(args.input)[0] + ".docx"
        
    if not os.path.exists(args.input):
        print(f"Error: File {args.input} not found.")
        sys.exit(1)
        
    markdown_to_docx(args.input, args.output)

if __name__ == "__main__":
    main()
